import streamlit as st
import PyPDF2
import nltk
from collections import Counter
from docx import Document
import difflib  # For calculating similarity in plagiarism check
from dotenv import load_dotenv
import base64
import os
from PIL import Image
import pdf2image
import google.generativeai as genai
from io import BytesIO
from fpdf import FPDF
import plotly.graph_objects as go
from auth import check_authentication
from docx.shared import Pt

# Set Streamlit page config at the top
st.set_page_config(page_title="JobFit AI - Smart Career Matcher", page_icon="🔍")

nltk.download('punkt')
nltk.download('stopwords')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Configure Google Generative AI
load_dotenv()
genai.configure(api_key=os.getenv("GENAI_API_KEY"))

check_authentication()

def get_gemini_response(resume_text, job_desc_text, prompt):
    """Fetches a response from Gemini API."""
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        input_text = f"Resume:\n{resume_text}\n\nJob Description:\n{job_desc_text}\n\nPrompt:\n{prompt}"
        response = model.generate_content(input_text)
        return response.text
    except Exception as e:
        st.error(f"Error in Gemini API: {e}")
        return None

def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    text = ''
    for page in reader.pages:
        text += page.extract_text()
    return text

def input_pdf_setup(pdf_file):
    return [extract_text_from_pdf(pdf_file)]

# Streamlit UI
st.title("**JobFit.AI System**")
st.subheader("About")
st.write("This sophisticated Applicant Tracking System, developed with Gemini Pro and Streamlit, seamlessly incorporates advanced features including resume match percentage, keyword analysis to identify missing criteria, and the generation of comprehensive profile summaries, enhancing the efficiency and precision of the candidate evaluation process for discerning talent acquisition professionals.")

# Sidebar for input
st.sidebar.header("Upload Your Job Description")
job_desc_file = st.sidebar.file_uploader("Upload Job Description (PDF)", type="pdf")


#Prompts
input_prompt1 = """
 You are an experienced Technical Human Resource Manager. Your task is to review the provided resume against the job description.
Please share your professional evaluation on whether the candidate's profile aligns with the role. Highlight the strengths and weaknesses
of the applicant in relation to the specified job requirements.
"""

input_prompt3 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of ATS functionality.
Your task is to evaluate the resume_pdf_content against the provided pdf_content and provide a match percentage.
The output should be a numerical percentage value only, without any additional text or symbols (e.g., 75).
"""


input_prompt4 = """
You are an skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality,
your task is to evaluate the resume against the provided job description. give me the relevant skills if the resume matches
the job description. The output should come as text containing all relevant skills required for given job description .
"""

input_prompt5 = """
You are an skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality,
your task is to evaluate the resume against the provided job description. give me the non-relevant skills if the resume matches
the job description. The output should come as text containing all non-relevant skills mentioned in resume that are not required for given job description .
"""

input_prompt7 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
Your task is to evaluate the resume against the provided job description and return List only the Relevant Projects with their required track according to job description, for the given job description.
The output should come as text containing all relevant projects required for given job description.
"""

input_prompt8 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
Your task is to evaluate the resume against the provided job description and return only the Recommended Skills required that are not available in resume but given in job description.
The output should come as text containing all recommended skills required for given job description.
"""

input_prompt9 = "Analyze the resume for grammatical errors, awkward phrasing, and clarity. Underline corrections and give suggestions for making it more ats friendly."

input_prompt10 = """
You are a professional career counselor and expert cover letter writer. Analyze the resume and job description 
to craft a personalized, compelling cover letter that:
1. Highlights the most relevant skills and experiences
2. Directly addresses key requirements in the job description
3. Demonstrates enthusiasm for the specific role
4. Uses a professional and engaging tone
5. Follows standard cover letter formatting and best practices
"""

# If a job description is uploaded
if job_desc_file is not None:
    pdf_content = input_pdf_setup(job_desc_file)
    job_desc_text = pdf_content[0]
    
    st.subheader("Your Resume")
    resume_file = st.file_uploader("Upload Your Resume (PDF)", type="pdf")
    
    if resume_file is not None:
        opt = st.sidebar.selectbox("Available Options", [
            "Choose an option", "Percentage match", "Show Relevant Skills", 
            "Non-relevant Skills", "Recommended Skills", "Relevant Projects", 
            "Tell Me About the Resume", "Grammatical Evaluation", "Generate Cover Letter"
        ])
        
        resume_pdf_content = input_pdf_setup(resume_file)
        resume_text = resume_pdf_content[0]
        
        if opt == "Percentage match":
            response = get_gemini_response(resume_text, job_desc_text, input_prompt3)
            st.subheader("Percentage Match")
            st.progress(int(response))
            st.write(f"Match: {response}%")
        
        elif opt == "Show Relevant Skills":
            relevant_skills = get_gemini_response(resume_text, job_desc_text, input_prompt4)
            st.write("Relevant Skills:")
            st.write(relevant_skills)
        
        elif opt == "Non-relevant Skills":
            non_relevant_skills = get_gemini_response(resume_text, job_desc_text, input_prompt5)
            st.write("Non-Relevant Skills:")
            st.write(non_relevant_skills)
        
        elif opt == "Relevant Projects":
            relevant_projects = get_gemini_response(resume_text, job_desc_text, input_prompt7)
            st.write("Relevant Projects:")
            st.write(relevant_projects)
        
        elif opt == "Recommended Skills":
            recommended_skills = get_gemini_response(resume_text, job_desc_text, input_prompt8)
            st.write("Recommended Skills:")
            st.write(recommended_skills)
        
        elif opt == "Tell Me About the Resume":
            evaluation_response = get_gemini_response(resume_text, job_desc_text, input_prompt1)
            if evaluation_response:
                st.subheader("Detailed Evaluation of Resume")
                st.write(evaluation_response)
        
        elif opt == "Grammatical Evaluation":
            response = get_gemini_response(resume_text, job_desc_text, input_prompt9)
            st.subheader("Grammatical Evaluation")
            st.write(response)
        
        elif opt == "Generate Cover Letter":
            st.subheader("Personalized Cover Letter")
            cover_letter_response = get_gemini_response(resume_text, job_desc_text, input_prompt10)
            if cover_letter_response:
                st.write(cover_letter_response)
                
                # TXT Download
                st.download_button(
                    label="Download Cover Letter as TXT",
                    data=cover_letter_response,
                    file_name="personalized_cover_letter.txt",
                    mime="text/plain"
                )
                
                # DOCX Download
                docx_buffer = BytesIO()
                doc = Document()
                doc.add_paragraph(cover_letter_response)
                for paragraph in doc.paragraphs:
                    paragraph.style = doc.styles.add_style('Custom', 1)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                st.download_button(
                    label="Download Cover Letter as DOCX",
                    data=docx_buffer,
                    file_name="personalized_cover_letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # PDF Download
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                for line in cover_letter_response.split('\n'):
                    pdf.cell(200, 10, txt=line, ln=True)
                pdf_buffer = BytesIO()
                pdf.output(pdf_buffer)
                pdf_buffer.seek(0)
                st.download_button(
                    label="Download Cover Letter as PDF",
                    data=pdf_buffer,
                    file_name="personalized_cover_letter.pdf",
                    mime="application/pdf"
                )
